# knowledge_loader.py
import os, re, json
from pathlib import Path
from bs4 import BeautifulSoup

# PDF
try:
    import pypdf
    HAS_PDF = True
except Exception:
    HAS_PDF = False

# DOCX
try:
    import docx  # package: python-docx
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

TEXT_EXT = {".txt", ".md", ".csv"}
HTML_EXT = {".html", ".htm"}
PDF_EXT  = {".pdf"}
DOCX_EXT = {".docx"}  # NEW

def _normalize(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()

def _read_text(fp: Path) -> str:
    try:
        return fp.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return fp.read_text(errors="ignore")

def _read_html(fp: Path) -> str:
    raw = _read_text(fp)
    soup = BeautifulSoup(raw, "html.parser")
    for tag in soup(["script","style","iframe","noscript"]):
        tag.decompose()
    return _normalize(soup.get_text(" "))

def _read_pdf(fp: Path) -> str:
    if not HAS_PDF:
        return ""
    try:
        reader = pypdf.PdfReader(str(fp))
        pages = [p.extract_text() or "" for p in reader.pages]
        return _normalize("\n".join(pages))
    except Exception:
        return ""

def _read_docx(fp: Path) -> str:
    if not HAS_DOCX:
        return ""
    try:
        d = docx.Document(str(fp))
        chunks = []
        # paragrafi
        for p in d.paragraphs:
            if p.text:
                chunks.append(p.text)
        # tabelle (righe > celle in pipe)
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text for c in row.cells]
                if any(cells):
                    chunks.append(" | ".join(cells))
        return _normalize("\n".join(chunks))
    except Exception:
        return ""

def walk_knowledge(base_dir: str) -> list[dict]:
    """
    Ritorna: [{path, relpath, mtime, text}, ...] per tutti i file supportati,
    scendendo ricorsivamente nelle sottocartelle.
    """
    base = Path(base_dir).resolve()
    out = []
    for fp in base.rglob("*"):
        if not fp.is_file():
            continue
        ext = fp.suffix.lower()
        text = ""
        if ext in TEXT_EXT:
            text = _read_text(fp)
        elif ext in HTML_EXT:
            text = _read_html(fp)
        elif ext in PDF_EXT:
            text = _read_pdf(fp)
        elif ext in DOCX_EXT:  # NEW
            text = _read_docx(fp)
        if not text:
            continue
        out.append({
            "path": str(fp),
            "relpath": str(fp.relative_to(base)),
            "mtime": fp.stat().st_mtime,
            "text": _normalize(text)
        })
    return out

def load_with_cache(base_dir: str, cache_path: str = ".knowledge_cache.json") -> list[dict]:
    # cache leggera su disco (rigenera solo se cambiano i file)
    try:
        cache = json.loads(Path(cache_path).read_text())
    except Exception:
        cache = {}
    recs = walk_knowledge(base_dir)
    changed = False
    for r in recs:
        key = r["path"]
        meta = {"mtime": r["mtime"], "size": Path(r["path"]).stat().st_size}
        if key in cache and cache[key].get("meta")==meta and cache[key].get("text"):
            r["text"] = cache[key]["text"]
        else:
            changed = True
            cache[key] = {"meta": meta, "text": r["text"], "relpath": r["relpath"]}
    if changed:
        try:
            Path(cache_path).write_text(json.dumps(cache, ensure_ascii=False))
        except Exception:
            pass
    return recs
